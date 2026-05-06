"""Tests for DOCX export."""

from __future__ import annotations

from pathlib import Path

import pytest

pytest.importorskip("docx")

from docx import Document

from inkline import export_docx


def test_export_docx_creates_file(tmp_path):
    out = tmp_path / "report.docx"
    result = export_docx("# Test Report\n\nBody text.", output_path=out, brand="minimal")
    assert result.exists()
    assert result.suffix == ".docx"


def test_export_docx_sets_core_title_from_heading(tmp_path):
    out = tmp_path / "title.docx"
    export_docx("# Extracted Title\n\nParagraph.", output_path=out, brand="minimal")
    doc = Document(out)
    assert doc.core_properties.title == "Extracted Title"


def test_export_docx_preserves_headings_and_lists(tmp_path):
    md = "# Main Title\n\n## Section One\n\n- Alpha\n- Beta\n\n1. First\n2. Second\n"
    out = tmp_path / "structure.docx"
    export_docx(md, output_path=out, brand="minimal")
    doc = Document(out)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "Main Title" in texts
    assert "Section One" in texts
    assert "Alpha" in texts
    assert "Second" in texts


def test_export_docx_creates_tables(tmp_path):
    md = (
        "# Table Test\n\n"
        "| Name | Value |\n"
        "| --- | --- |\n"
        "| Revenue | 10 |\n"
        "| Cost | 4 |\n"
    )
    out = tmp_path / "table.docx"
    export_docx(md, output_path=out, brand="minimal")
    doc = Document(out)
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "Name"
    assert doc.tables[0].cell(1, 0).text == "Revenue"
